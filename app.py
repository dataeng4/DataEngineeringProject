"""
app.py — Streamlit Frontend for Configuration-Driven SQLite Data Generation
============================================================================
Sections:
  1. Live Dashboard  — real-time record counter + key metrics
  2. Data Entry Form — YAML-driven dynamic form with duplicate detection
  3. Search & Export — filter, dataframe view, and FPDF biodata download
"""

import io
import textwrap
from datetime import datetime

import pandas as pd
import streamlit as st
import yaml
from fpdf import FPDF

import datetime

import re
from io import BytesIO
from docx import Document

# ── Local modules (provided by the user) ────────────────────────────────────
from database import check_duplicate, get_engine  # noqa: E402  (project-local)

# ============================================================
# Page Configuration
# ============================================================
st.set_page_config(
    page_title="UserForge · Data Studio",
    page_icon="🧬",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ============================================================
# Inline CSS — industrial-minimal dark theme
# ============================================================
st.markdown(
    """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Mono:wght@400;600&family=IBM+Plex+Sans:wght@300;400;600&display=swap');

    html, body, [class*="css"] {
        font-family: 'IBM Plex Sans', sans-serif;
        background-color: #0d0f12;
        color: #e2e8f0;
    }
    h1, h2, h3, h4 {
        font-family: 'IBM Plex Mono', monospace;
        letter-spacing: -0.03em;
    }
    /* Tab bar */
    div[data-baseweb="tab-list"] {
        gap: 0;
        border-bottom: 1px solid #2d3748;
    }
    button[data-baseweb="tab"] {
        font-family: 'IBM Plex Mono', monospace;
        font-size: 0.8rem;
        letter-spacing: 0.08em;
        text-transform: uppercase;
        padding: 10px 28px;
        border-radius: 0;
        color: #718096;
    }
    button[data-baseweb="tab"][aria-selected="true"] {
        color: #63b3ed;
        border-bottom: 2px solid #63b3ed;
        background: transparent;
    }
    /* Metric cards */
    div[data-testid="metric-container"] {
        background: #161b22;
        border: 1px solid #2d3748;
        border-radius: 8px;
        padding: 18px 24px;
    }
    div[data-testid="metric-container"] label {
        font-family: 'IBM Plex Mono', monospace;
        font-size: 0.72rem;
        letter-spacing: 0.1em;
        text-transform: uppercase;
        color: #718096;
    }
    div[data-testid="metric-container"] div[data-testid="stMetricValue"] {
        font-family: 'IBM Plex Mono', monospace;
        font-size: 2rem;
        font-weight: 600;
        color: #63b3ed;
    }
    /* Buttons */
    .stButton > button {
        font-family: 'IBM Plex Mono', monospace;
        font-size: 0.78rem;
        letter-spacing: 0.06em;
        text-transform: uppercase;
        background: #1a202c;
        border: 1px solid #4a5568;
        color: #e2e8f0;
        border-radius: 6px;
        transition: all 0.2s;
    }
    .stButton > button:hover {
        background: #2d3748;
        border-color: #63b3ed;
        color: #63b3ed;
    }
    /* Form submit button accent */
    .stFormSubmitButton > button {
        background: #1e3a5f !important;
        border-color: #63b3ed !important;
        color: #63b3ed !important;
        width: 100%;
    }
    /* Divider */
    hr { border-color: #2d3748; }
    /* Dataframe */
    .stDataFrame { border: 1px solid #2d3748; border-radius: 8px; }
    /* Input / select */
    input, textarea, select,
    div[data-baseweb="input"] input,
    div[data-baseweb="select"] {
        background-color: #161b22 !important;
        border-color: #2d3748 !important;
        color: #e2e8f0 !important;
        font-family: 'IBM Plex Mono', monospace !important;
        font-size: 0.85rem !important;
    }
    /* Header strip */
    .header-strip {
        background: linear-gradient(90deg, #0d0f12 0%, #161b22 60%, #1a2233 100%);
        border-bottom: 1px solid #2d3748;
        padding: 18px 32px 14px;
        margin-bottom: 8px;
    }
    .header-strip h1 { font-size: 1.6rem; color: #e2e8f0; margin: 0; }
    .header-strip span { color: #63b3ed; }
    .badge {
        display: inline-block;
        font-family: 'IBM Plex Mono', monospace;
        font-size: 0.65rem;
        letter-spacing: 0.12em;
        text-transform: uppercase;
        background: #1e3a5f;
        color: #63b3ed;
        border: 1px solid #2b6cb0;
        border-radius: 4px;
        padding: 2px 8px;
        margin-left: 10px;
        vertical-align: middle;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# ============================================================
# Master YAML Config (inline mirror — swap with yaml.safe_load
# of your config file if preferred)
# ============================================================
MASTER_YAML = """
pipeline:
  database: data/generated.db
  seed: 42
  batch_size: 5000

tables:
  - name: user_profiles
    target_records: 100000
    fields:
      - name: first_name
        sql_type: VARCHAR(50)
        label: First Name
        widget: text_input
        nullable: false
      - name: last_name
        sql_type: VARCHAR(50)
        label: Last Name
        widget: text_input
        nullable: false
      - name: age
        sql_type: INTEGER
        label: Age
        widget: number_input
        min: 18
        max: 75
        nullable: false
      - name: email
        sql_type: VARCHAR(120)
        label: Email Address
        widget: text_input
        unique: true
        nullable: false
      - name: mobile
        sql_type: VARCHAR(20)
        label: Mobile Number
        widget: text_input
        unique: true
        nullable: false
      - name: expertise
        sql_type: VARCHAR(100)
        label: Expertise / Job Title
        widget: text_input
        nullable: false
      - name: profile_type
        sql_type: VARCHAR(20)
        label: Profile Type
        widget: selectbox
        options:
          - admin
          - standard
          - premium
          - guest
        nullable: false
"""

config = yaml.safe_load(MASTER_YAML)
table_cfg = next(t for t in config["tables"] if t["name"] == "user_profiles")
fields = table_cfg["fields"]

# ============================================================
# DB Helpers
# ============================================================

@st.cache_resource
def get_conn():
    """Return a SQLAlchemy engine (cached for the session)."""
    return get_engine()


def fetch_metrics(engine):
    """Return (total_records, admin_count, premium_count, newest_email)."""
    with engine.connect() as conn:
        total = conn.execute(
            __import__("sqlalchemy").text("SELECT COUNT(*) FROM user_profiles")
        ).scalar()
        admins = conn.execute(
            __import__("sqlalchemy").text(
                "SELECT COUNT(*) FROM user_profiles WHERE profile_type='admin'"
            )
        ).scalar()
        premium = conn.execute(
            __import__("sqlalchemy").text(
                "SELECT COUNT(*) FROM user_profiles WHERE profile_type='premium'"
            )
        ).scalar()
    return int(total), int(admins), int(premium)


def insert_user(engine, data: dict):
    """Insert a single user record."""
    import sqlalchemy as sa

    table = sa.Table(
        "user_profiles",
        sa.MetaData(),
        autoload_with=engine,
    )
    with engine.begin() as conn:
        conn.execute(table.insert(), data)


def search_users(engine, query: str) -> pd.DataFrame:
    """Full-text-ish search across first_name, last_name, expertise."""
    import sqlalchemy as sa

    q = f"%{query}%"
    sql = sa.text(
        """
        SELECT id, first_name, last_name, age, email, mobile, expertise, profile_type
        FROM user_profiles
        WHERE first_name  LIKE :q
           OR last_name   LIKE :q
           OR expertise   LIKE :q
        LIMIT 200
        """
    )
    with engine.connect() as conn:
        rows = conn.execute(sql, {"q": q}).fetchall()
    cols = ["ID", "First Name", "Last Name", "Age", "Email", "Mobile", "Expertise", "Profile Type"]
    return pd.DataFrame(rows, columns=cols)


# ============================================================
# PDF Biodata Generator
# ============================================================

def generate_biodata_pdf(row: pd.Series) -> bytes:
    """Return PDF bytes for a single user-profile row."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # ── Header bar ──────────────────────────────────────────
    pdf.set_fill_color(13, 27, 42)
    pdf.rect(0, 0, 210, 38, "F")
    pdf.set_font("Helvetica", "B", 22)
    pdf.set_text_color(99, 179, 237)
    pdf.set_xy(14, 8)
    pdf.cell(0, 10, "UserForge · Profile Biodata", ln=True)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(160, 174, 192)
    pdf.set_x(14)
    pdf.cell(0, 6, f"Generated: {datetime.datetime.now().strftime('%d %b %Y  %H:%M:%S')}", ln=True)

    # ── Profile type badge ───────────────────────────────────
    badge_colors = {
        "admin": (220, 38, 38),
        "premium": (217, 119, 6),
        "standard": (37, 99, 235),
        "guest": (75, 85, 99),
    }
    ptype = str(row.get("Profile Type", "")).lower()
    bc = badge_colors.get(ptype, (75, 85, 99))
    pdf.set_fill_color(*bc)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 8)
    pdf.set_xy(155, 12)
    pdf.cell(40, 8, f"  {ptype.upper()}  ", border=0, fill=True, align="C")

    # ── Section: Personal Info ───────────────────────────────
    pdf.set_xy(14, 48)
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(40, 60, 80)
    pdf.cell(0, 7, "PERSONAL INFORMATION", ln=True)
    pdf.set_draw_color(99, 179, 237)
    pdf.set_line_width(0.5)
    pdf.line(14, pdf.get_y(), 196, pdf.get_y())
    pdf.ln(4)

    label_w = 55
    detail_fields = [
        ("Full Name",    f"{row['First Name']} {row['Last Name']}"),
        ("Age",          str(row["Age"])),
        ("Email",        str(row["Email"])),
        ("Mobile",       str(row["Mobile"])),
        ("Expertise",    str(row["Expertise"])),
        ("Profile Type", str(row["Profile Type"]).capitalize()),
        ("Record ID",    str(row["ID"])),
    ]

    for label, value in detail_fields:
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(80, 100, 120)
        pdf.set_x(14)
        pdf.cell(label_w, 9, label, border=0)
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(30, 40, 55)
        # Wrap long values
        wrapped = textwrap.shorten(value, width=55, placeholder="…")
        pdf.cell(0, 9, wrapped, ln=True)

    # ── Footer ────────────────────────────────────────────────
    pdf.set_y(-20)
    pdf.set_draw_color(200, 210, 220)
    pdf.set_line_width(0.3)
    pdf.line(14, pdf.get_y(), 196, pdf.get_y())
    pdf.set_font("Helvetica", "I", 7)
    pdf.set_text_color(160, 174, 192)
    pdf.set_x(14)
    pdf.cell(0, 8, "Confidential · Generated by UserForge Data Studio · For internal use only", align="C")

    return pdf.output(dest="S").encode("latin-1")



def generate_word(profile_data):
    doc = Document()
    doc.add_heading(f"Biodata: {profile_data['first_name']} {profile_data['last_name']}", level=1)
    
    for key, value in profile_data.items():
        doc.add_paragraph(f"{str(key).title().replace('_', ' ')}: {value}")
        
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# ============================================================
# App Header
# ============================================================
st.markdown(
    """
    <div class="header-strip">
        <h1>🧬 UserForge <span>·</span> Data Studio
            <span class="badge">v1.0.0</span>
        </h1>
    </div>
    """,
    unsafe_allow_html=True,
)

# ============================================================
# Tabs
# ============================================================
tab_dash, tab_form, tab_search = st.tabs(
    ["⬡  Live Dashboard", "✦  Data Entry", "⌕  Search & Export"]
)

engine = get_conn()

# ──────────────────────────────────────────────────────────────────────────────
# TAB 1 — LIVE DASHBOARD
# ──────────────────────────────────────────────────────────────────────────────
with tab_dash:
    st.markdown("#### Real-Time Database Metrics")
    st.caption("Counts refresh on every interaction — click **⟳ Refresh** to force an update.")

    if st.button("⟳  Refresh Metrics", key="refresh_dash"):
        st.cache_data.clear()

    total, admins, premium = fetch_metrics(engine)
    guest_approx = total - admins - premium  # rough remaining

    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Total Profiles", f"{total:,}", help="All rows in user_profiles")
    col2.metric("Admin Accounts", f"{admins:,}")
    col3.metric("Premium Accounts", f"{premium:,}")
    col4.metric("Other (est.)", f"{guest_approx:,}")

    st.divider()
    st.markdown("#### Profile Type Distribution")

    try:
        import sqlalchemy as sa

        with engine.connect() as conn:
            dist = conn.execute(
                sa.text(
                    "SELECT profile_type, COUNT(*) as cnt "
                    "FROM user_profiles GROUP BY profile_type"
                )
            ).fetchall()
        dist_df = pd.DataFrame(dist, columns=["Profile Type", "Count"])
        st.bar_chart(dist_df.set_index("Profile Type"), use_container_width=True, height=260)
    except Exception as exc:
        st.warning(f"Could not load distribution chart: {exc}")

    st.divider()
    st.markdown("#### Latest 10 Records")
    try:
        import sqlalchemy as sa

        with engine.connect() as conn:
            latest = conn.execute(
                sa.text(
                    "SELECT id, first_name, last_name, email, profile_type "
                    "FROM user_profiles ORDER BY id DESC LIMIT 10"
                )
            ).fetchall()
        latest_df = pd.DataFrame(
            latest, columns=["ID", "First Name", "Last Name", "Email", "Profile Type"]
        )
        st.dataframe(latest_df, use_container_width=True, hide_index=True)
    except Exception as exc:
        st.warning(f"Could not load latest records: {exc}")


# ──────────────────────────────────────────────────────────────────────────────
# TAB 2 — DATA ENTRY FORM (YAML-driven)
# ──────────────────────────────────────────────────────────────────────────────
with tab_form:
    st.markdown("#### New Profile Entry")
    st.caption("Form fields are generated dynamically from the master YAML configuration.")

    with st.form("user_entry_form", clear_on_submit=True):
        collected: dict = {}
        cols_left, cols_right = st.columns(2)
        col_toggle = True  # alternate columns for a tidy two-column layout

        for field in fields:
            fname = field["name"]
            label = field.get("label", fname.replace("_", " ").title())
            widget = field.get("widget", "text_input")
            nullable = field.get("nullable", True)

            # Choose column
            target_col = cols_left if col_toggle else cols_right
            col_toggle = not col_toggle

            with target_col:
                if widget == "text_input":
                    val = st.text_input(
                        label,
                        placeholder=f"Enter {label.lower()}…",
                        key=f"form_{fname}",
                    )
                    collected[fname] = val.strip() if val else None

                elif widget == "number_input":
                    val = st.number_input(
                        label,
                        min_value=field.get("min", 0),
                        max_value=field.get("max", 999),
                        step=1,
                        key=f"form_{fname}",
                    )
                    collected[fname] = int(val)

                elif widget == "selectbox":
                    options = field.get("options", [])
                    val = st.selectbox(label, options=options, key=f"form_{fname}")
                    collected[fname] = val

        st.markdown("")
        submitted = st.form_submit_button("➕  Submit Profile", use_container_width=True)

    if submitted:
        # ── Validation ────────────────────────────────────────
        missing = [
            field.get("label", f["name"])
            for field in fields
            if not field.get("nullable", True)
            for f in [field]
            if not collected.get(field["name"])
            and field.get("widget", "text_input") == "text_input"
        ]
        if missing:
            st.error(f"Required field(s) missing: {', '.join(missing)}")
        else:
            email = collected.get("email", "")
            mobile = collected.get("mobile", "")

            # ── Strict Regex Validation ───────────────────────────
            email_pattern = r"^[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+$"
            mobile_pattern = r"^\d{10}$" # Forces exactly 10 numeric digits

            if not re.match(email_pattern, email):
                st.error("⚠️ Invalid Email Format! Please include an '@' and domain.")
            elif not re.match(mobile_pattern, mobile):
                st.error("⚠️ Invalid Mobile! Please enter exactly 10 digits without text.")
            elif check_duplicate(email, mobile):
                st.error("⚠️ Email or Mobile already exists in the database!")
            else:
                try:
                    collected["created_at"] = datetime.datetime.now()
                    insert_user(engine, collected)
                    total_new, _, _ = fetch_metrics(engine)
                    st.success(
                        f"✅  Profile created successfully! "
                        f"Total records: **{total_new:,}**"
                    )
                    st.balloons()
                except Exception as exc:
                    st.error(f"Database error: {exc}")

    # ── Info expander: YAML field map ─────────────────────────
    with st.expander("🗂  View YAML Field Configuration"):
        st.code(MASTER_YAML, language="yaml")


# ──────────────────────────────────────────────────────────────────────────────
# TAB 3 — SEARCH & EXPORT
# ──────────────────────────────────────────────────────────────────────────────
with tab_search:
    st.markdown("#### Search Profiles")
    st.caption("Filter by first name, last name, or expertise. Up to 200 results shown.")

    search_query = st.text_input(
        "Search",
        placeholder="e.g.  'Alice', 'Engineer', 'Smith' …",
        label_visibility="collapsed",
        key="search_input",
    )

    if search_query:
        with st.spinner("Querying database…"):
            results_df = search_users(engine, search_query)

        if results_df.empty:
            st.info("No profiles matched your query.")
        else:
            st.markdown(
                f"<small style='color:#718096;font-family:IBM Plex Mono,monospace'>"
                f"{len(results_df)} result(s) found</small>",
                unsafe_allow_html=True,
            )
            st.dataframe(results_df, use_container_width=True, hide_index=True)

            st.divider()
            st.markdown("#### Download Biodata")
            st.caption("Click on a profile below to view details and download documents.")

            # ── Create a beautiful drop-down for EVERY searched user ──
            for index, row in results_df.iterrows():
                with st.expander(f"👤  {row['First Name']} {row['Last Name']} - {row['Expertise']}"):
                    
                    # --- Keep your awesome custom UI ---
                    pc1, pc2 = st.columns([1, 2])
                    with pc1:
                        badge_color_map = {
                            "admin": "#DC2626", "premium": "#D97706",
                            "standard": "#2563EB", "guest": "#4B5563",
                        }
                        pt = str(row.get("Profile Type", "")).lower()
                        bc_hex = badge_color_map.get(pt, "#4B5563")
                        st.markdown(
                            f"""
                            <div style='background:#161b22;border:1px solid #2d3748;border-radius:10px;padding:20px 18px;'>
                                <div style='font-family:IBM Plex Mono,monospace;font-size:1.5rem;font-weight:600;color:#e2e8f0;margin-bottom:4px'>
                                    {row['First Name']} {row['Last Name']}
                                </div>
                                <div style='font-size:0.8rem;color:#718096;margin-bottom:12px'>
                                    ID #{row['ID']}
                                </div>
                                <span style='background:{bc_hex};color:#fff;border-radius:4px;font-family:IBM Plex Mono,monospace;font-size:0.68rem;letter-spacing:0.1em;padding:3px 10px;text-transform:uppercase'>
                                    {pt}
                                </span>
                            </div>
                            """, unsafe_allow_html=True
                        )
                    with pc2:
                        info_rows = [
                            ("Age", row["Age"]), ("Email", row["Email"]),
                            ("Mobile", row["Mobile"]), ("Expertise", row["Expertise"])
                        ]
                        for lbl, val in info_rows:
                            st.markdown(
                                f"<div style='font-family:IBM Plex Mono,monospace;font-size:0.78rem;color:#718096;margin-top:8px'>{lbl}</div>"
                                f"<div style='font-size:0.95rem;color:#e2e8f0'>{val}</div>",
                                unsafe_allow_html=True
                            )
                    
                    st.write("") # Small spacer
                    btn1, btn2 = st.columns(2)
                    fname_safe = f"{row['First Name']}_{row['Last Name']}".replace(" ", "_").lower()
                    
                    # --- PDF Download Button ---
                    with btn1:
                        try:
                            pdf_bytes = generate_biodata_pdf(row)
                            st.download_button(
                                label="📄 Download PDF",
                                data=pdf_bytes,
                                file_name=f"biodata_{fname_safe}_{row['ID']}.pdf",
                                mime="application/pdf",
                                use_container_width=True,
                                key=f"pdf_{index}" # Unique key required for loops
                            )
                        except Exception as exc:
                            st.error(f"PDF failed: {exc}")

                    # --- Word Download Button ---
                    with btn2:
                        try:
                            word_bytes = generate_word(row.to_dict())
                            st.download_button(
                                label="📝 Download Word",
                                data=word_bytes,
                                file_name=f"biodata_{fname_safe}_{row['ID']}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                                key=f"word_{index}" # Unique key required for loops
                            )
                        except Exception as exc:
                            st.error(f"Word failed: {exc}")

    else:
        st.markdown(
            "<div style='text-align:center;padding:60px 0;"
            "font-family:IBM Plex Mono,monospace;color:#4a5568;font-size:0.9rem'>"
            "↑ Type a name or expertise above to begin searching</div>",
            unsafe_allow_html=True,
        )